#!/usr/bin/env python3
"""
Pipeline completo: CSV bruto → canais endêmicos → canais por faixa etária → index.html

Uso:
python pipeline.py dados_ids.csv --pop 210000 --output index.html

Este script:
1. Roda compute_channels.py para gerar channel_data.json (58 canais)
2. Agrega dados por faixa etária para 20 agravos prioritários
3. Computa canais Gamma-Poisson por faixa etária
4. Gera boletim enriquecido
5. Monta o index.html final com todos os dados embutidos
6. (Opcional) Gera boletim epidemiológico em DOCX (--boletim)
"""

import json
import os
import sys
import subprocess
import argparse
from datetime import datetime
from pathlib import Path
import numpy as np
import pandas as pd
import io

# Imports para boletim DOCX (step 6)
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════
# Config
# ══════════════════════════════════════════════════════════════════════

MC_SAMPLES = 30_000
QUANTILES = [0.10, 0.25, 0.50, 0.75, 0.90]
RNG_SEED = 2026
MAX_SE = 52
TRAINING_YEARS_DEFAULT = ['2023', '2024', '2025']  # 2021/2022 excluídos (implantação)
MONITOR_YEAR = '2026'
FALLBACK_SHAPE = 0.1
FALLBACK_RATE = 1.0

AGE_GROUPS = {
    "Bebês (0-1)":          (0, 1),
    "Crianças (2-11)":      (2, 11),
    "Adolescentes (12-17)": (12, 17),
    "Adultos jovens (18-39)":(18, 39),
    "Meia-idade (40-59)":   (40, 59),
    "Idosos (60+)":         (60, 200),
}

KEY_AGRAVOS_AGE = [
    "Todos os atendimentos",
    "X - Aparelho respiratório",
    "SINAN: Influenza NE",
    "SINAN: Dengue",
    "SINAN: Diarréia/gastroenterite",
    "SINAN: Pneumonia NE",
    "Sínd. Gripal (IVAS+Febre)",
    "Sínd. Respiratória (J09-J22)",
    "Gastroenterites (A09+K52)",
    "Saúde Mental",
    "Dor Osteomuscular",
    "Cardiovascular Aguda",
    "XIII - Sistema osteomuscular",
    "V - Transtornos mentais",
    "XIX - Lesões e causas externas",
    "XVIII - Sintomas e sinais",
    "Acidentes de Trânsito",
    "Agressões",
    "Tentativas de Suicídio",
    "Dermatológica Aguda",
]

DESC_TO_CID = {}
SYNDROME_DEFS = {}

# ══════════════════════════════════════════════════════════════════════
# Gamma-Poisson (simplificado para faixas etárias)
# ══════════════════════════════════════════════════════════════════════

from math import lgamma, log, exp, inf

def estimate_params_mom_simple(cases_arr):
    x = np.array(cases_arr, dtype=float)
    m = np.mean(x)
    v = np.var(x, ddof=1) if len(x) > 1 else m + 1
    if m <= 0:
        return FALLBACK_SHAPE, FALLBACK_RATE
    if v <= m:
        v = m + 0.01
    denom = v - m
    if denom <= 0:
        denom = 0.01
    a_hat = max(m * m / denom, 0.01)
    b_hat = max(m / denom, 0.001)
    return a_hat, b_hat

def nb_loglik(x_arr, shape, rate):
    r, p = shape, rate / (rate + 1.0)
    ll = 0.0
    for x in x_arr:
        x = int(x)
        ll += lgamma(x + r) - lgamma(r) - lgamma(x + 1)
        ll += r * log(p + 1e-300) + x * log(1 - p + 1e-300)
    return ll

def estimate_params_mle_simple(cases_arr, a0=None, b0=None):
    x = np.array(cases_arr, dtype=float)
    if a0 is None or b0 is None:
        a0, b0 = estimate_params_mom_simple(cases_arr)
    best_a, best_b, best_ll = a0, b0, -inf
    for scale in [2.0, 0.5, 0.1]:
        a_grid = np.exp(np.linspace(log(max(best_a * exp(-scale), 0.001)),
                                    log(best_a * exp(scale)), 15))
        b_grid = np.exp(np.linspace(log(max(best_b * exp(-scale), 0.0001)),
                                    log(best_b * exp(scale)), 15))
        for a in a_grid:
            for b in b_grid:
                ll = nb_loglik(x, a, b)
                if ll > best_ll:
                    best_ll, best_a, best_b = ll, a, b
    return best_a, best_b

def mc_quantiles(shape, rate, n_samples=MC_SAMPLES, rng=None):
    if rng is None:
        rng = np.random.default_rng(RNG_SEED)
    lam = rng.gamma(shape, 1.0 / rate, size=n_samples)
    x = rng.poisson(lam)
    return [int(np.quantile(x, q)) for q in QUANTILES]

def classify_zone(value, thresholds):
    p10, p25, p50, p75, p90 = thresholds
    if value <= p25:   return 'sucesso'
    elif value <= p50: return 'seguranca'
    elif value <= p75: return 'alerta'
    elif value <= p90: return 'epidemico'
    else:              return 'emergencia'

# ══════════════════════════════════════════════════════════════════════
# Step 1: Rodar compute_channels.py
# ══════════════════════════════════════════════════════════════════════

def step1_compute_channels(csv_path, pop, channel_json='channel_data.json'):
    print("\n" + "=" * 60)
    print("STEP 1: Computando canais endêmicos principais")
    print("=" * 60)
    cmd = [
        sys.executable, 'compute_channels.py',
        csv_path,
        '--pop', str(pop),
        '--output', channel_json,
        '--agravos', 'all',
    ]
    result = subprocess.run(cmd, capture_output=False)
    if result.returncode != 0:
        print("ERRO: compute_channels.py falhou!")
        sys.exit(1)
    with open(channel_json, 'r') as f:
        data = json.load(f)
    print(f"  → {len(data['channels'])} canais gerados")
    return data

# ══════════════════════════════════════════════════════════════════════
# Step 2: Agregar dados por faixa etária
# ══════════════════════════════════════════════════════════════════════

def step2_age_group_data(csv_path, channel_data):
    print("\n" + "=" * 60)
    print("STEP 2: Agregando dados por faixa etária")
    print("=" * 60)

    import importlib.util
    spec = importlib.util.spec_from_file_location("compute_channels", "compute_channels.py")
    cc = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(cc)

    for enc in ['utf-8', 'latin-1', 'cp1252']:
        try:
            df = pd.read_csv(csv_path, sep=';', encoding=enc)
            break
        except (UnicodeDecodeError, pd.errors.ParserError):
            continue

    cols = {c.lower().strip(): c for c in df.columns}
    col_cid = cols.get('cid_descricao', cols.get('cid', 'cid_descricao'))
    has_faixa_etaria = 'faixa_etaria' in cols
    has_idade        = 'idade' in cols
    has_ano_epi      = 'ano_epi' in cols
    has_data         = 'data' in cols

    if has_ano_epi and 'semana_epi' in cols:
        print("  → Detectado formato GAS agregado (ano_epi + semana_epi)")
        df['ano_epi'] = df[cols['ano_epi']]
        df['se_epi']  = df[cols['semana_epi']]
    elif has_data:
        col_date = cols['data']
        df[col_date] = pd.to_datetime(df[col_date], dayfirst=True, errors='coerce')
        df = df.dropna(subset=[col_date])
        if hasattr(cc, 'epi_week'):
            epi = df[col_date].apply(cc.epi_week)
            df['ano_epi'] = epi.apply(lambda x: x[0])
            df['se_epi']  = epi.apply(lambda x: x[1])
        else:
            from datetime import timedelta as _td
            def _epi_week_fallback(dt):
                dow = dt.isoweekday() % 7
                sun = dt - _td(days=dow)
                wed = sun + _td(days=3)
                ano = wed.year
                jan1 = pd.Timestamp(ano, 1, 1)
                j1d = jan1.isoweekday() % 7
                if j1d <= 3:
                    fs = jan1 - _td(days=j1d)
                else:
                    fs = jan1 + _td(days=7 - j1d)
                se = (sun - fs).days // 7 + 1
                return (ano, max(1, se))
            epi_fb = df[col_date].apply(_epi_week_fallback)
            df['ano_epi'] = epi_fb.apply(lambda x: x[0])
            df['se_epi']  = epi_fb.apply(lambda x: x[1])
    else:
        print("  ⚠ Sem coluna 'data' nem 'ano_epi' — skip age channels")
        return {}

    col_qty = 'quantidade'
    if col_qty not in cols:
        df['quantidade'] = 1
    else:
        col_qty = cols['quantidade']

    if 'cid_codigo' in cols:
        df['cid_code'] = df[cols['cid_codigo']].astype(str).str.strip().str.upper()
        df.loc[df['cid_code'].isin(['', 'NAN', 'NONE']), 'cid_code'] = pd.NA
        print(f"  → Coluna cid_codigo encontrada: {df['cid_code'].notna().mean():.0%} com código")
    else:
        desc_to_cid = cc.DESC_TO_CID if hasattr(cc, 'DESC_TO_CID') else {}
        df['cid_code'] = df[col_cid].astype(str).str.strip().str.upper().map(desc_to_cid)

    FAIXA_TO_AGE_GROUP = {
        '<1': 'Bebês (0-1)', '1-4': 'Bebês (0-1)',
        '5-9': 'Crianças (2-11)', '10-14': 'Crianças (2-11)',
        '15-19': 'Adolescentes (12-17)',
        '20-29': 'Adultos jovens (18-39)', '30-39': 'Adultos jovens (18-39)',
        '40-49': 'Meia-idade (40-59)',     '50-59': 'Meia-idade (40-59)',
        '60-69': 'Idosos (60+)', '70-79': 'Idosos (60+)', '80+': 'Idosos (60+)',
        'NI': None,
    }

    sinan_reverse = {}
    if hasattr(cc, 'SINAN_MAP'):
        for code, disease in cc.SINAN_MAP.items():
            sinan_reverse.setdefault(disease, set()).add(code)

    if 'cid_code' in df.columns:
        df['_sinan']   = df['cid_code'].apply(lambda x: cc.cid_to_sinan(x)   if hasattr(cc, 'cid_to_sinan')   else 'Outros')
        df['_chapter'] = df['cid_code'].apply(lambda x: cc.cid_to_chapter(x) if hasattr(cc, 'cid_to_chapter') else None)
    else:
        df['_sinan']   = 'Outros'
        df['_chapter'] = None

    if has_faixa_etaria:
        col_age = cols['faixa_etaria']
        print("  → Classificando por faixa_etaria (string)")
        df['faixa'] = df[col_age].astype(str).str.strip().map(FAIXA_TO_AGE_GROUP)
    elif has_idade:
        col_age = cols['idade']
        print("  → Classificando por idade (numérica)")
        def parse_age(val):
            try:
                v = str(val).strip()
                num = ''.join(c for c in v.split()[0] if c.isdigit() or c == '.')
                return int(float(num)) if num else None
            except:
                return None
        df['idade_num'] = df[col_age].apply(parse_age)
        df = df.dropna(subset=['idade_num'])
        df['idade_num'] = df['idade_num'].astype(int)
        def get_age_group(age):
            for name, (lo, hi) in AGE_GROUPS.items():
                if lo <= age <= hi:
                    return name
            return None
        df['faixa'] = df['idade_num'].apply(get_age_group)
    else:
        print("  ⚠ Nenhuma coluna de idade/faixa_etaria encontrada — skip age channels")
        return {}

    df = df.dropna(subset=['faixa'])

    available_agravos = list(channel_data['channels'].keys())
    selected = [a for a in KEY_AGRAVOS_AGE if a in available_agravos]
    age_data = {}

    for agravo in selected:
        age_data[agravo] = {}
        for faixa in AGE_GROUPS:
            df_f = df[df['faixa'] == faixa]
            if agravo == "Todos os atendimentos":
                df_a = df_f
            elif agravo.startswith("SINAN: "):
                sinan_disease = agravo[7:]
                df_a = df_f[df_f['_sinan'] == sinan_disease]
            elif any(agravo.startswith(f"{rom} -") for rom in
                     ['I','II','III','IV','V','VI','VII','VIII','IX','X','XI','XII',
                      'XIII','XIV','XV','XVI','XVII','XVIII','XIX','XX','XXI','XXII']):
                df_a = df_f[df_f['_chapter'] == agravo]
            elif hasattr(cc, 'SYNDROME_DEFS') and agravo in cc.SYNDROME_DEFS:
                syn_codes = cc.SYNDROME_DEFS[agravo]
                if 'cid_code' in df.columns:
                    mask = df_f['cid_code'].apply(
                        lambda x: any(str(x).startswith(c) for c in syn_codes)
                        if pd.notna(x) else False)
                    df_a = df_f[mask]
                else:
                    df_a = df_f
            else:
                cid_code_match = agravo.split(' - ')[0].strip() if ' - ' in agravo else agravo
                if 'cid_code' in df.columns:
                    df_a = df_f[df_f['cid_code'] == cid_code_match]
                else:
                    df_a = df_f

            if len(df_a) == 0:
                continue

            counts = df_a.groupby(['ano_epi', 'se_epi'])[col_qty].sum().reset_index(name='n')
            year_se = {}
            for _, row in counts.iterrows():
                yr = str(int(row['ano_epi']))
                se = str(int(row['se_epi']))
                if int(se) > MAX_SE:
                    continue
                if yr not in year_se:
                    year_se[yr] = {}
                year_se[yr][se] = int(row['n'])
            if year_se:
                age_data[agravo][faixa] = year_se

    print(f"  → {len(age_data)} agravos × faixas etárias agregados")
    return age_data

# ══════════════════════════════════════════════════════════════════════
# Step 3: Computar canais por faixa etária
# ══════════════════════════════════════════════════════════════════════

def step3_age_channels(age_data):
    print("\n" + "=" * 60)
    print("STEP 3: Computando canais por faixa etária")
    print("=" * 60)

    rng = np.random.default_rng(RNG_SEED)
    results = {}
    total = sum(len(age_data[a]) for a in age_data)
    done = 0

    for agravo in sorted(age_data.keys()):
        results[agravo] = {}
        for age_group in sorted(age_data[agravo].keys()):
            done += 1
            grp = age_data[agravo][age_group]
            channels = {}
            raw = {}
            classifications = {}
            all_years = sorted(set(TRAINING_YEARS_DEFAULT + [MONITOR_YEAR]))
            for yr in all_years:
                raw[yr] = {}
                if yr in grp:
                    for se_s, count in grp[yr].items():
                        se = int(se_s)
                        if 1 <= se <= MAX_SE:
                            raw[yr][se] = int(count)

            for se in range(1, MAX_SE + 1):
                train = [int(grp.get(yr, {}).get(str(se), 0)) for yr in TRAINING_YEARS_DEFAULT]
                if sum(train) == 0:
                    thresholds = [0, 0, 0, 1, 2]
                else:
                    a0, b0 = estimate_params_mom_simple(train)
                    try:
                        a, b = estimate_params_mle_simple(train, a0, b0)
                    except:
                        a, b = a0, b0
                    thresholds = mc_quantiles(a, b, n_samples=MC_SAMPLES, rng=rng)
                channels[se] = {
                    'p10': thresholds[0], 'p25': thresholds[1],
                    'p50': thresholds[2], 'p75': thresholds[3], 'p90': thresholds[4]
                }

            for yr in all_years:
                classifications[yr] = {}
                for se in range(1, MAX_SE + 1):
                    val = raw.get(yr, {}).get(se, 0)
                    if se in channels:
                        th = [channels[se]['p10'], channels[se]['p25'],
                              channels[se]['p50'], channels[se]['p75'], channels[se]['p90']]
                        classifications[yr][str(se)] = classify_zone(val, th)

            raw_str      = {yr: {str(k): v for k, v in raw[yr].items()} for yr in raw}
            channels_str = {str(k): v for k, v in channels.items()}
            results[agravo][age_group] = {
                'channels':       channels_str,
                'raw':            raw_str,
                'classifications': classifications
            }
            if done % 20 == 0 or done == total:
                print(f"  [{done}/{total}] {agravo} / {age_group}")

    print(f"  → {total} combinações computadas")
    return results

# ══════════════════════════════════════════════════════════════════════
# Step 4: Gerar boletim enriquecido
# ══════════════════════════════════════════════════════════════════════

def step4_boletim(channel_data):
    print("\n" + "=" * 60)
    print("STEP 4: Gerando boletim enriquecido")
    print("=" * 60)

    channels = channel_data['channels']
    priority_agravos = [
        ("SINAN: Dengue",                   "ALTA"),
        ("SINAN: Diarréia/gastroenterite",  "ALTA"),
        ("SINAN: Pneumonia NE",             "ALTA"),
        ("Todos os atendimentos",           "MODERADA"),
        ("I - Doenças infecciosas e parasitárias", "ALTA"),
        ("X - Aparelho respiratório",       "MODERADA"),
        ("XVIII - Sintomas e sinais",       "MODERADA"),
        ("XIII - Sistema osteomuscular",    "MODERADA"),
        ("SINAN: Influenza NE",             "MODERADA"),
        ("V - Transtornos mentais",         "MODERADA"),
        ("IX - Aparelho circulatório",      "BAIXA"),
        ("XIV - Aparelho geniturinário",    "BAIXA"),
        ("XII - Pele e tecido subcutâneo",  "BAIXA"),
        ("SINAN: Sífilis NE",              "BAIXA"),
    ]

    def find_channel(name, channels):
        if name in channels:
            return name, channels[name]
        name_upper = name.upper()
        for key in channels:
            if name_upper in key.upper() or key.upper() in name_upper:
                return key, channels[key]
        if name.startswith("SINAN: "):
            short = name[7:]
            for key in channels:
                if short.upper() in key.upper():
                    return key, channels[key]
        return None, None

    boletim = []
    for name, prio in priority_agravos:
        matched_name, ch = find_channel(name, channels)
        if not ch:
            continue
        name = matched_name

        se_list = ch['se_list']
        raw     = ch.get('raw', [])
        years   = ch['years']

        total_2025 = sum(r.get('c2025', 0) for r in raw)
        total_2026 = sum(r.get('c2026', 0) for r in raw)
        hist_years  = [y for y in years if 2022 <= y <= 2024]
        hist_totals = [sum(r.get(f'c{y}', 0) for r in raw) for y in hist_years]
        media_hist  = int(np.mean(hist_totals)) if hist_totals else 0
        var_pct     = round((total_2025 - media_hist) / max(media_hist, 1) * 100, 1)

        vals_2025 = [(r.get('c2025', 0), se) for r, se in zip(raw, se_list)]
        pico_val, pico_se = max(vals_2025, key=lambda x: x[0])

        cls_2025 = ch.get('classifications', {}).get('2025', [])
        se_p90   = sum(1 for z in cls_2025 if z == 'emergencia')

        last_se = 0
        for r, se in zip(raw, se_list):
            if r.get('c2026', 0) > 0:
                last_se = se

        zone_counts_2025 = {'sucesso': 0, 'seguranca': 0, 'alerta': 0, 'epidemico': 0, 'emergencia': 0}
        for z in cls_2025:
            if z in zone_counts_2025:
                zone_counts_2025[z] += 1

        cls_2026 = ch.get('classifications', {}).get('2026', [])
        zone_counts_2026 = {'sucesso': 0, 'seguranca': 0, 'alerta': 0, 'epidemico': 0, 'emergencia': 0}
        for z in cls_2026:
            if z in zone_counts_2026:
                zone_counts_2026[z] += 1

        if var_pct > 10:
            tend = f"Aumento de {var_pct}% em 2025 vs média 2022-2024."
        elif var_pct < -10:
            tend = f"Redução de {abs(var_pct)}% em 2025 vs média 2022-2024."
        else:
            tend = f"Estável de {var_pct}% em 2025 vs média 2022-2024."

        acao = "Manter vigilância ativa." if prio in ("ALTA", "MODERADA") else "Monitoramento de rotina."

        ultima_zona = cls_2026[last_se - 1] if last_se > 0 and last_se <= len(cls_2026) else 'sem dados'
        obs_ult     = raw[last_se - 1].get('c2026', 0) if last_se > 0 else 0
        ch_2026     = ch.get('channels', {}).get('2026', [])
        p90_ult     = ch_2026[last_se - 1][4] if ch_2026 and last_se > 0 else 1
        p50_ult     = ch_2026[last_se - 1][2] if ch_2026 and last_se > 0 else 1

        boletim.append({
            'name':              name,
            'prioridade':        prio,
            'tendencia':         tend,
            'sazonalidade':      f"Pico na SE {pico_se}.",
            'acao':              acao,
            'total_2025':        total_2025,
            'se_p90_2025':       se_p90,
            'total_2026':        total_2026,
            'se_2026':           last_se,
            'variacao_pct':      var_pct,
            'media_hist':        media_hist,
            'classificacao_2025': zone_counts_2025,
            'classificacao_2026': zone_counts_2026,
            'pico_val_2025':     pico_val,
            'pico_se_2025':      pico_se,
            'pico_val_2026':     max((r.get('c2026', 0) for r in raw), default=0),
            'pico_se_2026':      max(((r.get('c2026', 0), se) for r, se in zip(raw, se_list)),
                                    key=lambda x: x[0], default=(0, 1))[1],
            'ultima_se_zona':    ultima_zona,
            'ultima_se_obs':     obs_ult,
            'ultima_se_p90':     p90_ult,
            'ultima_se_p50':     p50_ult,
            'media_semanal_2026': round(total_2026 / max(last_se, 1), 1),
            'media_semanal_2025': round(total_2025 / 52, 1),
        })

    print(f"  → {len(boletim)} agravos no boletim")
    return boletim

# ══════════════════════════════════════════════════════════════════════
# Step 5: Montar HTML final
# ══════════════════════════════════════════════════════════════════════

def step5_generate_html(channel_data, age_data, age_channels, boletim,
                        template_html, output_html):
    print("\n" + "=" * 60)
    print("STEP 5: Gerando HTML final")
    print("=" * 60)

    with open(template_html, 'r') as f:
        html = f.read()

    now    = datetime.now()
    now_br = now.strftime('%d/%m/%Y')

    ch_total = channel_data['channels'].get('Total de atendimentos', {})
    raw      = ch_total.get('raw', [])
    se_list  = ch_total.get('se_list', [])
    last_se  = 0
    for r, se in zip(raw, se_list):
        if r.get('c2026', 0) > 0:
            last_se = se

    data_json = json.dumps(channel_data, separators=(',', ':'), ensure_ascii=False)
    import re
    html = re.sub(r'const DATA = \{.*?\};\s*\n', '', html, count=1, flags=re.DOTALL)
    script_pos = html.find('<script type="text/babel">') + len('<script type="text/babel">')
    html = html[:script_pos] + f"\nconst DATA = {data_json};\n" + html[script_pos:]

    age_json = json.dumps(age_data, separators=(',', ':'), ensure_ascii=False)
    html = re.sub(r'const AGE_GROUP_DATA = \{.*?\};\s*\n', '', html, count=1, flags=re.DOTALL)
    insert_after = html.find('const DATA = ')
    insert_after = html.find(';\n', insert_after) + 2
    html = html[:insert_after] + f"const AGE_GROUP_DATA = {age_json};\n" + html[insert_after:]

    ac_compact = build_age_channels_compact(age_channels)
    ac_json = json.dumps(ac_compact, separators=(',', ':'), ensure_ascii=False)
    html = re.sub(r'const AGE_CHANNELS = \{.*?\};\s*\n', '', html, count=1, flags=re.DOTALL)
    insert_after = html.find('const AGE_COLORS = ')
    insert_after = html.find('};\n', insert_after) + 3
    html = html[:insert_after] + f"const AGE_CHANNELS = {ac_json};\n" + html[insert_after:]

    bol_json = json.dumps(boletim, separators=(',', ':'), ensure_ascii=False)
    html = re.sub(r'const BOLETIM_DATA = \[.*?\];\s*\n', '', html, count=1, flags=re.DOTALL)
    insert_after = html.find('const AGE_CHANNELS = ')
    insert_after = html.find(';\n', insert_after) + 2
    html = html[:insert_after] + f"const BOLETIM_DATA = {bol_json};\n" + html[insert_after:]

    html = re.sub(
        r'Última extração: <b[^>]*>[\d/]+</b>',
        f'Última extração: <b style={{{{ color: "#fbbf24" }}}}>{now_br}</b>', html)
    html = re.sub(
        r'Dados até: <b[^>]*>[\d/]+</b> \(SE \d+\)',
        f'Dados até: <b style={{{{ color: "#60a5fa" }}}}>{now_br}</b> (SE {last_se})', html)
    html = re.sub(
        r'Gerado em: <b[^>]*>[\d/]+</b>',
        f'Gerado em: <b style={{{{ color: "#a5b4fc" }}}}>{now_br}</b>', html)
    html = re.sub(r'SE \d+/2026', f'SE {last_se}/2026', html)

    with open(output_html, 'w') as f:
        f.write(html)
    size_mb = os.path.getsize(output_html) / (1024 * 1024)
    print(f"  → {output_html}: {size_mb:.1f} MB")
    print(f"  → Datas atualizadas: extração={now_br}, SE={last_se}/2026")


def build_age_channels_compact(ac_data):
    result = {}
    for agravo in ac_data:
        result[agravo] = {}
        for age_group in ac_data[agravo]:
            agd = ac_data[agravo][age_group]
            ch  = agd['channels']
            raw = agd['raw']
            years = sorted(raw.keys(), key=lambda x: int(x))
            channels_compact = {}
            for yr in years:
                yr_ch = []
                for se in range(1, 53):
                    se_s = str(se)
                    c = ch.get(se_s, {'p10':0,'p25':0,'p50':0,'p75':0,'p90':0})
                    yr_ch.append([c['p10'],c['p25'],c['p50'],c['p75'],c['p90']])
                channels_compact[yr] = yr_ch
            raw_compact = []
            for se in range(1, 53):
                entry = {}
                for yr in years:
                    entry[f'c{yr}'] = raw.get(yr, {}).get(str(se), 0)
                raw_compact.append(entry)
            result[agravo][age_group] = {
                'years':    [int(y) for y in years],
                'se_list':  list(range(1, 53)),
                'channels': channels_compact,
                'raw':      raw_compact,
            }
    return result

# ══════════════════════════════════════════════════════════════════════
# Helpers para Step 6
# ══════════════════════════════════════════════════════════════════════

def _zona_label(zona):
    MAP = {
        'sucesso':    ('SUCESSO',    '1A7942'),
        'seguranca':  ('SEGURANÇA',  '2E75B6'),
        'alerta':     ('ALERTA',     'F5A623'),
        'epidemico':  ('EPIDÊMICO',  'D62828'),
        'emergencia': ('EMERGÊNCIA', '7B0D8E'),
        'sem dados':  ('SEM DADOS',  'AAAAAA'),
    }
    return MAP.get(zona, ('?', 'AAAAAA'))


def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _gerar_grafico_canal(channel_data, agravos_plot, se_atual, monitor_year='2026'):
    channels = channel_data.get('channels', {})
    CORES = {
        'Todos os atendimentos':          '#2E75B6',
        'X - Aparelho respiratório':      '#E05C2E',
        'Sínd. Respiratória (J09-J22)':   '#E07B2E',
        'SINAN: Dengue':                  '#D62828',
        'Gastroenterites (A09+K52)':      '#1A7942',
        'SINAN: Diarréia/gastroenterite': '#1A9942',
        'Sínd. Gripal (IVAS+Febre)':      '#9B59B6',
    }
    disponiveis = [a for a in agravos_plot if a in channels]
    if not disponiveis:
        disponiveis = list(channels.keys())[:3]

    n    = len(disponiveis)
    fig, axes = plt.subplots(n, 1, figsize=(10, 3.2 * n), squeeze=False)
    fig.patch.set_facecolor('#F8F9FA')

    for idx, agravo in enumerate(disponiveis):
        ax = axes[idx][0]
        ch = channels[agravo]
        raw_list = ch.get('raw', [])
        se_list  = ch.get('se_list', [])
        years    = ch.get('years', [])
        ch_data  = ch.get('channels', {})
        clf_data = ch.get('classifications', {})
        yr_str   = str(monitor_year)

        canal = ch_data.get(yr_str)
        if canal is None and ch_data:
            canal = ch_data[max(ch_data.keys())]

        obs_2026 = []
        ses_plot = []
        for i, se in enumerate(se_list):
            if se > se_atual:
                break
            val = raw_list[i].get(f'c{monitor_year}', 0) if i < len(raw_list) else 0
            obs_2026.append(val)
            ses_plot.append(se)

        train_years = [y for y in years if y != int(monitor_year) and y >= 2023]
        hist_med = []
        for i, se in enumerate(ses_plot):
            vals = [raw_list[i].get(f'c{y}', 0) for y in train_years if i < len(raw_list)]
            hist_med.append(sum(vals) / max(len(vals), 1))

        if canal and ses_plot:
            p25 = [canal[i][1] if i < len(canal) else 0 for i in range(len(ses_plot))]
            p50 = [canal[i][2] if i < len(canal) else 0 for i in range(len(ses_plot))]
            p75 = [canal[i][3] if i < len(canal) else 0 for i in range(len(ses_plot))]
            p90 = [canal[i][4] if i < len(canal) else 0 for i in range(len(ses_plot))]
            ax.fill_between(ses_plot, 0,    p25, alpha=0.18, color='#1A7942', label='Sucesso (≤P25)')
            ax.fill_between(ses_plot, p25,  p50, alpha=0.18, color='#2E75B6', label='Segurança (P25–P50)')
            ax.fill_between(ses_plot, p50,  p75, alpha=0.18, color='#F5A623', label='Alerta (P50–P75)')
            ax.fill_between(ses_plot, p75,  p90, alpha=0.25, color='#D62828', label='Epidêmico (P75–P90)')
            ax.fill_between(ses_plot, p90,  [max(p90)*1.5 or 1]*len(ses_plot),
                            alpha=0.15, color='#7B0D8E', label='Emergência (>P90)')
            ax.plot(ses_plot, p50, '--', color='#555', linewidth=0.8, alpha=0.6)

        if hist_med:
            ax.plot(ses_plot, hist_med, ':', color='#888', linewidth=1.0,
                    label=f'Média {min(train_years) if train_years else "hist."}–{max(train_years) if train_years else ""}')

        cor = CORES.get(agravo, '#333333')
        if obs_2026:
            ax.plot(ses_plot, obs_2026, '-o', color=cor, linewidth=2,
                    markersize=4, label=f'Observado {monitor_year}', zorder=5)
            clf_yr = clf_data.get(yr_str, [])
            for i, (se, val) in enumerate(zip(ses_plot, obs_2026)):
                zone = clf_yr[i] if i < len(clf_yr) else 'sucesso'
                _, hx = _zona_label(zone)
                r = int(hx[0:2], 16)/255
                g = int(hx[2:4], 16)/255
                b = int(hx[4:6], 16)/255
                ax.plot(se, val, 'o', color=(r,g,b), markersize=6, zorder=6)

        ax.axvline(x=se_atual, color='#333', linewidth=1.2, linestyle='--', alpha=0.5)
        ylim = ax.get_ylim()
        ax.text(se_atual + 0.3, ylim[1] * 0.9 if ylim[1] > 0 else 1,
                f'SE {se_atual}', fontsize=7, color='#333')

        nome_curto = agravo.replace('SINAN: ', '').strip()
        if len(nome_curto) > 40:
            nome_curto = nome_curto[:38] + '…'
        ax.set_title(nome_curto, fontsize=9, fontweight='bold', loc='left', pad=4)
        ax.set_xlabel('Semana Epidemiológica', fontsize=8)
        ax.set_ylabel('Atendimentos', fontsize=8)
        ax.set_xlim(1, max(se_list) if se_list else 52)
        ax.tick_params(labelsize=7)
        ax.set_facecolor('#FAFAFA')
        ax.grid(True, alpha=0.3, linewidth=0.5)
        ax.legend(loc='upper left', fontsize=6, ncol=2, framealpha=0.8)

    fig.suptitle(
        f'Curvas Sindromicas — Rio Claro/SP | SE {se_atual}/{monitor_year}',
        fontsize=10, fontweight='bold', y=1.01
    )
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=130, bbox_inches='tight',
                facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf.read()

# ══════════════════════════════════════════════════════════════════════
# Step 6: Boletim DOCX com 4 itens da curva sindromica
# ══════════════════════════════════════════════════════════════════════

def step6_boletim_docx(boletim_data, channel_data, se_num, output_html,
                       municipio='Rio Claro/SP', monitor_year='2026',
                       se_label=None):
    print("\n" + "=" * 60)
    print("STEP 6: Gerando boletim DOCX")
    print("=" * 60)

    doc     = DocxDocument()
    now     = datetime.now()
    nome_arq = f'boletim_SE{se_num}.docx'
    channels = channel_data.get('channels', {})

    # ── Página A4 ────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    def add_spacer():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    def add_heading(text, level=1, color_hex='1F4E79'):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(
                int(color_hex[0:2], 16),
                int(color_hex[2:4], 16),
                int(color_hex[4:6], 16))
        return p

    def _get_zona_obs(nome_orig):
        ch = channels.get(nome_orig)
        if not ch:
            for k in channels:
                if nome_orig.upper() in k.upper():
                    ch = channels[k]
                    break
        if not ch:
            return 'sem dados', 0, 0, 0
        raw_list = ch.get('raw', [])
        clf      = ch.get('classifications', {}).get(str(monitor_year), [])
        ch_yr    = ch.get('channels', {}).get(str(monitor_year)) or \
                   (ch.get('channels', {}).get(max(ch.get('channels', {}).keys(), default=monitor_year)))
        se_idx   = se_num - 1
        if se_idx < 0 or se_idx >= len(raw_list):
            return 'sem dados', 0, 0, 0
        obs  = raw_list[se_idx].get(f'c{monitor_year}', 0)
        zona = clf[se_idx] if se_idx < len(clf) else 'sem dados'
        p50  = ch_yr[se_idx][2] if ch_yr and se_idx < len(ch_yr) else 0
        p90  = ch_yr[se_idx][4] if ch_yr and se_idx < len(ch_yr) else 0
        return zona, obs, p50, p90

    # ════════════════════════════════════════════════════════════════
    # SEÇÃO 1 — Cabeçalho + situação da SE corrente
    # ════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BOLETIM EPIDEMIOLÓGICO SEMANAL')
    run.bold = True; run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'Fundação Municipal de Saúde — {municipio}')
    r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _se_header = se_label if se_label else f'Semana Epidemiológica {se_num}/{monitor_year}'
    r3 = p3.add_run(
        f'{_se_header} '
        f'· Gerado em {now.strftime("%d/%m/%Y")} '
        f'· EpiKanalis / EpiKinesis Inteligência Ltda.')
    r3.font.size = Pt(9); r3.italic = True
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Linha divisória
    bp = doc.add_paragraph()
    pPr  = bp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F4E79')
    pBdr.append(bot); pPr.append(pBdr)
    add_spacer()

    add_heading(f'1. Situação na SE {se_num}/{monitor_year}', level=1)

    AGRAVOS_DESTAQUE = [
        'Todos os atendimentos', 'X - Aparelho respiratório',
        'SINAN: Dengue', 'Sínd. Gripal (IVAS+Febre)',
        'Gastroenterites (A09+K52)', 'V - Transtornos mentais',
    ]
    tbl_data = []
    for nome in AGRAVOS_DESTAQUE:
        zona, obs, p50, p90 = _get_zona_obs(nome)
        if obs == 0 and zona == 'sem dados':
            continue
        nome_curto = nome.replace('SINAN: ', '').replace('Sínd. ', 'Sínd. ')
        if len(nome_curto) > 38: nome_curto = nome_curto[:36] + '…'
        tbl_data.append((nome_curto, obs, p50, p90, zona))

    if tbl_data:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, col_name in enumerate([f'Agravo / Síndrome', f'Obs SE {se_num}', 'P50', 'P90', 'Situação']):
            hdr[i].text = col_name
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
                hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
            _set_cell_bg(hdr[i], '1F4E79')
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        for nome_curto, obs, p50, p90, zona in tbl_data:
            label, hex_c = _zona_label(zona)
            row = table.add_row().cells
            row[0].text = nome_curto
            row[1].text = str(obs)
            row[2].text = str(p50)
            row[3].text = str(p90)
            row[4].text = label
            _set_cell_bg(row[4], hex_c)
            if row[4].paragraphs[0].runs:
                row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
                row[4].paragraphs[0].runs[0].bold = True
            for cell in row:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
        widths = [Cm(7.5), Cm(2.2), Cm(1.8), Cm(1.8), Cm(2.8)]
        for rw in table.rows:
            for i, cell in enumerate(rw.cells):
                cell.width = widths[i]

    add_spacer()

    # ════════════════════════════════════════════════════════════════
    # SEÇÃO 2 — Curvas sindromicas com canal endêmico
    # ════════════════════════════════════════════════════════════════
    add_heading('2. Curvas Sindromicas e Canais Endêmicos', level=1)

    p_desc = doc.add_paragraph()
    r_desc = p_desc.add_run(
        f'Comportamento semanal dos principais agravos em {monitor_year} comparado '
        f'ao canal endêmico Gamma-Poisson (base: 2023–2025). Faixas de cor: '
        f'verde=sucesso, azul=segurança, amarelo=alerta, vermelho=epidêmico, roxo=emergência.')
    r_desc.font.size = Pt(10)
    add_spacer()

    AGRAVOS_GRAF = [
        'Todos os atendimentos', 'X - Aparelho respiratório',
        'SINAN: Dengue', 'Sínd. Gripal (IVAS+Febre)',
        'Gastroenterites (A09+K52)',
    ]
    disponiveis_graf = []
    for a in AGRAVOS_GRAF:
        if a in channels:
            disponiveis_graf.append(a)
        else:
            for k in channels:
                if a.upper() in k.upper():
                    disponiveis_graf.append(k); break
        if len(disponiveis_graf) >= 4:
            break

    try:
        png_bytes  = _gerar_grafico_canal(channel_data, disponiveis_graf, se_num, monitor_year)
        img_stream = io.BytesIO(png_bytes)
        doc.add_picture(img_stream, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(
            f'Figura 1. Curvas sindromicas + canal endêmico Gamma-Poisson — '
            f'SE 1–{se_num}/{monitor_year}. Base: UPAs Rio Claro (Maestro/IDS Saúde). '
            f'Método: EpiKanalis v2 (EpiKinesis Inteligência Ltda.)')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(8); run.italic = True
        print("  → Gráfico gerado e inserido")
    except Exception as e:
        doc.add_paragraph(f'[Gráfico não disponível: {e}]')
        print(f"  ⚠ Gráfico falhou: {e}")

    add_spacer()

    # ════════════════════════════════════════════════════════════════
    # SEÇÃO 3 — Tabela de classificação de zonas
    # ════════════════════════════════════════════════════════════════
    add_heading('3. Classificação de Zonas — Agravos Prioritários', level=1)

    p_desc2 = doc.add_paragraph()
    r_desc2 = p_desc2.add_run(
        f'Classificação bayesiana na SE {se_num}/{monitor_year} com observado vs limiares preditivos. '
        f'Tendência calculada sobre as 3 SEs anteriores (↑ >15%, ↓ <-15%, → estável).')
    r_desc2.font.size = Pt(10)
    add_spacer()

    AGRAVOS_TABELA = [
        ('SINAN: Dengue',                   'Alta'),
        ('SINAN: Diarréia/gastroenterite',  'Alta'),
        ('SINAN: Pneumonia NE',             'Alta'),
        ('Sínd. Gripal (IVAS+Febre)',        'Moderada'),
        ('Sínd. Respiratória (J09-J22)',     'Moderada'),
        ('X - Aparelho respiratório',        'Moderada'),
        ('Gastroenterites (A09+K52)',         'Moderada'),
        ('Todos os atendimentos',            'Baixa'),
        ('V - Transtornos mentais',          'Baixa'),
        ('XIII - Sistema osteomuscular',     'Baixa'),
        ('XIX - Lesões e causas externas',   'Baixa'),
        ('IX - Aparelho circulatório',       'Baixa'),
    ]

    table2 = doc.add_table(rows=1, cols=6)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    for i, col_name in enumerate(['Agravo', 'Prior.', f'Obs SE {se_num}', 'P25–P90', 'Situação', 'Tendência']):
        hdr2[i].text = col_name
        if hdr2[i].paragraphs[0].runs:
            hdr2[i].paragraphs[0].runs[0].bold = True
            hdr2[i].paragraphs[0].runs[0].font.size = Pt(8)
        _set_cell_bg(hdr2[i], '2E4057')
        if hdr2[i].paragraphs[0].runs:
            hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    for nome_orig, prio in AGRAVOS_TABELA:
        ch = channels.get(nome_orig)
        nome_real = nome_orig
        if not ch:
            for k in channels:
                if nome_orig.upper() in k.upper():
                    ch = channels[k]; nome_real = k; break
        if not ch:
            continue

        raw_list = ch.get('raw', [])
        clf      = ch.get('classifications', {}).get(str(monitor_year), [])
        ch_yr    = ch.get('channels', {}).get(str(monitor_year)) or \
                   (ch.get('channels', {}).get(max(ch.get('channels', {}).keys(), default=monitor_year)))
        se_idx   = se_num - 1
        if se_idx < 0 or se_idx >= len(raw_list):
            continue

        obs  = raw_list[se_idx].get(f'c{monitor_year}', 0)
        zona = clf[se_idx] if se_idx < len(clf) else 'sem dados'
        p25  = ch_yr[se_idx][1] if ch_yr and se_idx < len(ch_yr) else 0
        p90  = ch_yr[se_idx][4] if ch_yr and se_idx < len(ch_yr) else 0

        vals_ant = [raw_list[max(0, se_idx-j)].get(f'c{monitor_year}', 0)
                    for j in range(1,4) if se_idx-j >= 0]
        if vals_ant:
            media_ant = sum(vals_ant) / len(vals_ant)
            tend = '↑ Subindo' if obs > media_ant*1.15 else ('↓ Caindo' if obs < media_ant*0.85 else '→ Estável')
        else:
            tend = '—'

        label, hex_c = _zona_label(zona)
        nome_curto = nome_real.replace('SINAN: ','').replace('Sínd. ','S. ')
        if len(nome_curto) > 32: nome_curto = nome_curto[:30] + '…'

        row = table2.add_row().cells
        row[0].text = nome_curto
        row[1].text = prio
        row[2].text = str(obs)
        row[3].text = f'{p25}–{p90}'
        row[4].text = label
        row[5].text = tend
        _set_cell_bg(row[4], hex_c)
        if row[4].paragraphs[0].runs:
            row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            row[4].paragraphs[0].runs[0].bold = True
        for cell in row:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(8)

    w2 = [Cm(5.5), Cm(1.8), Cm(1.8), Cm(2.5), Cm(2.5), Cm(2.3)]
    for rw in table2.rows:
        for i, cell in enumerate(rw.cells):
            cell.width = w2[i]

    add_spacer()

    # ════════════════════════════════════════════════════════════════
    # SEÇÃO 4 — Análise interpretativa
    # ════════════════════════════════════════════════════════════════
    add_heading('4. Análise Interpretativa', level=1)

    zona_dengue, obs_dengue, _, _ = _get_zona_obs('SINAN: Dengue')
    zona_resp,   obs_resp,   _, _ = _get_zona_obs('X - Aparelho respiratório')
    zona_gripal, obs_gripal, _, _ = _get_zona_obs('Sínd. Gripal (IVAS+Febre)')
    zona_gastro, obs_gastro, _, _ = _get_zona_obs('Gastroenterites (A09+K52)')

    criticos = sum(1 for z in [zona_dengue, zona_resp, zona_gripal, zona_gastro]
                   if z in ('epidemico','emergencia'))
    alertas  = sum(1 for z in [zona_dengue, zona_resp, zona_gripal, zona_gastro]
                   if z == 'alerta')

    if criticos >= 2:
        nivel  = 'ATENÇÃO EPIDEMIOLÓGICA ELEVADA'; cor_n = 'D62828'
        intro  = (f'A SE {se_num}/{monitor_year} registra {criticos} agravos em zona epidêmica ou '
                  f'emergência. Vigilância ativa intensificada e avaliação de resposta recomendadas.')
    elif criticos == 1 or alertas >= 2:
        nivel  = 'SITUAÇÃO DE ATENÇÃO'; cor_n = 'F5A623'
        intro  = (f'A SE {se_num}/{monitor_year} apresenta elevação em pelo menos um agravo prioritário. '
                  f'Monitoramento reforçado e avaliação das condições de atendimento recomendados.')
    else:
        nivel  = 'SITUAÇÃO DENTRO DO ESPERADO'; cor_n = '1A7942'
        intro  = (f'A SE {se_num}/{monitor_year} apresenta comportamento dentro dos parâmetros '
                  f'históricos. Manter vigilância epidemiológica de rotina.')

    p_nivel = doc.add_paragraph()
    r_nivel = p_nivel.add_run(f'▶ AVALIAÇÃO GERAL: {nivel}')
    r_nivel.bold = True; r_nivel.font.size = Pt(11)
    r_nivel.font.color.rgb = RGBColor(int(cor_n[0:2],16), int(cor_n[2:4],16), int(cor_n[4:6],16))

    p_intro = doc.add_paragraph()
    p_intro.add_run(intro).font.size = Pt(10)
    add_spacer()

    sindromes = [
        (zona_dengue, obs_dengue, 'Dengue',
         'Acionar protocolo de intensificação vetorial.' if zona_dengue in ('epidemico','emergencia')
         else 'Monitoramento de rotina.'),
        (zona_resp, obs_resp, 'Síndrome respiratória',
         'Avaliar capacidade de atendimento nas UPAs.' if zona_resp in ('alerta','epidemico','emergencia')
         else 'Sem elevação relevante.'),
        (zona_gripal, obs_gripal, 'Síndrome gripal (IVAS + febre)',
         'Reforçar orientação sobre uso racional de antibióticos.' if zona_gripal in ('alerta','epidemico','emergencia')
         else 'Comportamento sazonal dentro do esperado.'),
        (zona_gastro, obs_gastro, 'Gastroenterites',
         'Monitorar possíveis fontes de contaminação hídrica ou alimentar.' if zona_gastro in ('epidemico','emergencia')
         else 'Sem sinais de surto.'),
    ]
    for zona, obs, nome, rec in sindromes:
        label, _ = _zona_label(zona)
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{nome}: {obs} atendimentos — {label}. {rec}').font.size = Pt(10)

    add_spacer()
    add_heading('Recomendações Operacionais', level=2, color_hex='2E4057')

    recs = []
    if zona_dengue in ('epidemico','emergencia'):
        recs.append('Acionar protocolo de intensificação do controle vetorial — dengue em zona crítica.')
    if zona_resp in ('alerta','epidemico','emergencia') or zona_gripal in ('alerta','epidemico','emergencia'):
        recs.append('Avaliar estoque de insumos respiratórios e escala nas UPAs.')
    if not recs:
        recs.append('Manter monitoramento semanal de rotina.')
    recs.append(f'Próximo boletim: SE {se_num+1}/{monitor_year}. '
                f'Sistema EpiKanalis (EpiKinesis Inteligência Ltda.).')
    for rec in recs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(rec).font.size = Pt(10)

    # Rodapé
    add_spacer()
    p_rod = doc.add_paragraph()
    r_rod = p_rod.add_run(
        f'Fonte: Sistema Maestro/IDS Saúde — UPAs Rio Claro (Cervezão e UPA 29). '
        f'Método: Canal Endêmico Gamma-Poisson Hierárquico (EpiKanalis v2). '
        f'Base 2023–2025 (2022 excluído). '
        f'Contato: EpiKinesis Inteligência Ltda. | ekokubun@rc.unesp.br')
    r_rod.font.size = Pt(7.5); r_rod.italic = True
    r_rod.font.color.rgb = RGBColor(0x88,0x88,0x88)
    p_rod.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(nome_arq)
    size_kb = os.path.getsize(nome_arq) / 1024
    print(f"  → Boletim salvo: {nome_arq} ({size_kb:.0f} KB)")
    return nome_arq

# ══════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description='Pipeline completo: CSV → Dashboard HTML atualizado')
    parser.add_argument('input', nargs='?', default=None,
                        help='CSV de entrada (dados brutos IDS Saúde). '
                             'Opcional quando --from-json é usado.')
    parser.add_argument('--pop',    type=int, default=210000,
                        help='População do município (default: 210000)')
    parser.add_argument('--output', '-o', default='index.html',
                        help='HTML de saída (default: index.html)')
    parser.add_argument('--template', default='index.html',
                        help='HTML template (default: index.html existente)')
    parser.add_argument('--boletim', action='store_true',
                        help='Gerar boletim epidemiológico em DOCX além do HTML')
    parser.add_argument('--se-num', type=int, default=None,
                        help='Número da SE corrente para o boletim (auto-detecta se omitido)')
    parser.add_argument('--se-year', type=int, default=None,
                        help='Ano da SE (default: 2026)')
    parser.add_argument('--se-label', type=str, default=None,
                        help='Rótulo textual da SE para o cabeçalho do boletim '
                             '(ex.: "SE 19/2026 — 6 a 12 de maio de 2026")')
    parser.add_argument('--from-json', type=str, default=None,
                        help='Pula steps 1-3 e 5 e gera APENAS o boletim DOCX a partir '
                             'de um channel_data.json existente. Implica --boletim.')

    args = parser.parse_args()

    # ── Modo boletim-only: a partir de channel_data.json ──────────────
    if args.from_json:
        print("╔══════════════════════════════════════════════════════════╗")
        print("║ Pipeline — Modo boletim DOCX (sem reprocessar pipeline)  ║")
        print("╚══════════════════════════════════════════════════════════╝")
        print(f"  channel_data.json: {args.from_json}")
        if args.se_label:
            print(f"  {args.se_label}")
        with open(args.from_json, 'r', encoding='utf-8') as f:
            channel_data = json.load(f)
        boletim = step4_boletim(channel_data)

        se_num = args.se_num
        if se_num is None:
            for item in boletim:
                if item.get('se_2026', 0) > 0:
                    se_num = item['se_2026']
                    break
        if se_num is None:
            from datetime import date
            hoje = date.today()
            jan1 = date(hoje.year, 1, 1)
            se_num = max(1, (hoje - jan1).days // 7 + 1)
            print(f"  SE auto-detectada: {se_num}")

        monitor_year = str(args.se_year) if args.se_year else '2026'
        step6_boletim_docx(
            boletim_data=boletim,
            channel_data=channel_data,
            se_num=se_num,
            output_html=args.output,
            monitor_year=monitor_year,
            se_label=args.se_label,
        )
        print(f"\n✓ Boletim gerado: boletim_SE{se_num}.docx")
        return

    # ── Modo completo: requer CSV ─────────────────────────────────────
    if args.input is None:
        parser.error("input (CSV) é obrigatório quando --from-json não é usado")

    print("╔══════════════════════════════════════════════════════════╗")
    print("║ Pipeline Canal Endêmico — Atualização Automática         ║")
    print("╚══════════════════════════════════════════════════════════╝")
    print(f"  CSV: {args.input}")
    print(f"  Pop: {args.pop}")
    print(f"  Output: {args.output}")
    if args.boletim:
        print(f"  Boletim DOCX: SIM (SE={args.se_num or 'auto-detectar'})")

    channel_data = step1_compute_channels(args.input, args.pop)
    age_data     = step2_age_group_data(args.input, channel_data)
    age_channels = step3_age_channels(age_data)
    boletim      = step4_boletim(channel_data)
    step5_generate_html(channel_data, age_data, age_channels, boletim,
                        args.template, args.output)

    if args.boletim:
        se_num = args.se_num
        if se_num is None:
            for item in boletim:
                if item.get('se_2026', 0) > 0:
                    se_num = item['se_2026']
                    break
        if se_num is None:
            from datetime import date
            hoje = date.today()
            jan1 = date(hoje.year, 1, 1)
            se_num = max(1, (hoje - jan1).days // 7 + 1)
            print(f"  SE auto-detectada: {se_num}")

        step6_boletim_docx(
            boletim_data=boletim,
            channel_data=channel_data,
            se_num=se_num,
            output_html=args.output
        )

    print("\n✓ Pipeline concluído com sucesso!")
    print(f"  Dashboard: {args.output}")
    if args.boletim:
        print(f"  Boletim DOCX: boletim_SE{se_num}.docx")


if __name__ == '__main__':
    main()
